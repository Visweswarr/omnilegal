from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
ASSET_DIR = REPORT_DIR / "detailed_design_report_assets"
SCREENSHOT_DIR = REPORT_DIR / "project_report_assets"
OUT = REPORT_DIR / "OMNILEGAL_DETAILED_DESIGN_THINKING_REPORT.docx"

ASSET_DIR.mkdir(parents=True, exist_ok=True)


COLORS = {
    "ink": "111111",
    "muted": "666666",
    "gold": "B7791F",
    "amber": "D97706",
    "green": "15803D",
    "red": "B91C1C",
    "blue": "1D4ED8",
    "paper": "F7F1E6",
    "line": "D9D2C3",
}


def rgb(hex_value: str) -> RGBColor:
    hex_value = hex_value.strip("#")
    return RGBColor(int(hex_value[0:2], 16), int(hex_value[2:4], 16), int(hex_value[4:6], 16))


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            pass
    return ImageFont.load_default()


def draw_multiline(draw, xy, text, fnt, fill, max_width, line_gap=8):
    words = text.split()
    lines = []
    line = ""
    for word in words:
        trial = f"{line} {word}".strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            line = trial
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap
    return y


def make_box_diagram(path: Path, title: str, subtitle: str, boxes: list[tuple[str, str]], footer: str = ""):
    w, h = 1800, 980
    img = Image.new("RGB", (w, h), "#FBF8F0")
    d = ImageDraw.Draw(img)
    title_font = font(56, True)
    sub_font = font(28)
    box_title = font(30, True)
    box_body = font(22)
    small = font(20)

    d.rectangle([0, 0, w, 120], fill=f"#{COLORS['ink']}")
    d.text((70, 34), title, font=title_font, fill="#F7F1E6")
    d.text((70, 132), subtitle, font=sub_font, fill=f"#{COLORS['muted']}")

    margin_x = 70
    gap = 28
    cols = min(3, len(boxes))
    box_w = (w - margin_x * 2 - gap * (cols - 1)) // cols
    box_h = 260
    y0 = 220
    for i, (head, body) in enumerate(boxes):
        row = i // cols
        col = i % cols
        x = margin_x + col * (box_w + gap)
        y = y0 + row * (box_h + 42)
        d.rounded_rectangle([x, y, x + box_w, y + box_h], radius=18, fill="#FFFFFF", outline=f"#{COLORS['line']}", width=3)
        d.rectangle([x, y, x + 12, y + box_h], fill=f"#{COLORS['gold']}")
        d.text((x + 34, y + 28), head, font=box_title, fill=f"#{COLORS['ink']}")
        draw_multiline(d, (x + 34, y + 78), body, box_body, f"#{COLORS['muted']}", box_w - 70, 8)
    if footer:
        d.text((70, h - 70), footer, font=small, fill=f"#{COLORS['muted']}")
    img.save(path)


def make_flow_diagram(path: Path):
    w, h = 1900, 760
    img = Image.new("RGB", (w, h), "#FBF8F0")
    d = ImageDraw.Draw(img)
    title_font = font(54, True)
    body_font = font(24)
    label_font = font(28, True)
    small_font = font(19)
    d.rectangle([0, 0, w, 112], fill=f"#{COLORS['ink']}")
    d.text((64, 30), "OmniLegal Technical Flow", font=title_font, fill="#F7F1E6")
    steps = [
        ("User Query", "A legal question, claim, clause, or topic"),
        ("API Router", "FastAPI routes select the correct workflow"),
        ("RAG Pipeline", "classify -> retrieve -> analyze -> synthesize"),
        ("Sources", "Local corpus, datasets, live registries"),
        ("Verifier", "Citation existence, quote match, support checks"),
        ("Answer UI", "Cited answer, map, audit, IRAC, or report"),
    ]
    x, y = 58, 205
    box_w, box_h, gap = 270, 220, 40
    for i, (head, body) in enumerate(steps):
        bx = x + i * (box_w + gap)
        d.rounded_rectangle([bx, y, bx + box_w, y + box_h], radius=18, fill="#FFFFFF", outline=f"#{COLORS['line']}", width=3)
        d.text((bx + 24, y + 28), head, font=label_font, fill=f"#{COLORS['ink']}")
        draw_multiline(d, (bx + 24, y + 78), body, body_font, f"#{COLORS['muted']}", box_w - 48, 8)
        if i < len(steps) - 1:
            ax1 = bx + box_w + 8
            ay = y + box_h // 2
            ax2 = bx + box_w + gap - 8
            d.line([ax1, ay, ax2, ay], fill=f"#{COLORS['gold']}", width=5)
            d.polygon([(ax2, ay), (ax2 - 18, ay - 12), (ax2 - 18, ay + 12)], fill=f"#{COLORS['gold']}")
    d.text((64, 560), "The central design rule: generated legal prose is never the final product until sources and citations are visible to the user.", font=small_font, fill=f"#{COLORS['muted']}")
    img.save(path)


def make_design_cycle(path: Path):
    boxes = [
        ("Empathize", "Understand pain points of students, researchers, MUN delegates, and non-law users."),
        ("Define", "Convert vague legal-research struggles into clear problem statements and design criteria."),
        ("Ideate", "Generate workflows such as Research Console, Citation Forensics, Conflict Atlas, and Comparative IRAC."),
        ("Prototype", "Build React, Chainlit, FastAPI, RAG retrieval, datasets, prompts, and verification flows."),
        ("Test", "Run smoke tests, citation checks, team walkthroughs, and UI feedback loops."),
    ]
    make_box_diagram(
        path,
        "Design Thinking Applied to OmniLegal",
        "The project is not only an AI pipeline; it is an iterative human-centered design process.",
        boxes,
        "Each stage produces concrete artifacts: personas, journey maps, wireframes, prototype screens, metrics, and next actions.",
    )


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9D2C3", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def format_table(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                    run.font.size = Pt(9.2)
            if header and r == 0:
                set_cell_shading(cell, COLORS["ink"])
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = rgb("FFFFFF")
                        run.font.bold = True
            elif r % 2 == 1:
                set_cell_shading(cell, "F8F5EE")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = value
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    format_table(table)
    doc.add_paragraph()
    return table


def add_h(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Georgia" if level <= 2 else "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
        run.font.color.rgb = rgb(COLORS["ink"] if level == 1 else COLORS["gold"])
    return p


def add_p(doc: Document, text: str, bold_start: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.12
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        r.font.color.rgb = rgb(COLORS["ink"])
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(10.2)
        if run.font.color.rgb is None:
            run.font.color.rgb = rgb("2D2D2D")
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.08
        p.add_run(item)
        for run in p.runs:
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            run.font.size = Pt(9.8)


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)
        for run in p.runs:
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            run.font.size = Pt(9.8)


def add_callout(doc: Document, title: str, body: str, fill="FFF7E6"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "E7C77E", "8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    r.font.color.rgb = rgb(COLORS["ink"])
    p2 = cell.add_paragraph()
    p2.add_run(body)
    for run in p2.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9.6)
        run.font.color.rgb = rgb("333333")
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, width_inches: float = 6.5):
    if not path.exists():
        add_callout(doc, "Missing figure", f"The expected image was not found: {path}", "FDECEC")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = "Arial"
    r.font.size = Pt(8.8)
    r.font.color.rgb = rgb(COLORS["muted"])


def set_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = rgb("2D2D2D")

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Georgia" if style_name != "Heading 3" else "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.color.rgb = rgb(COLORS["ink"] if style_name == "Heading 1" else COLORS["gold"])
        style.font.bold = True


def add_cover(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("OMNILEGAL AI")
    r.font.name = "Georgia"
    r.font.size = Pt(32)
    r.font.bold = True
    r.font.color.rgb = rgb(COLORS["ink"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Detailed Project and Design Thinking Report")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.color.rgb = rgb(COLORS["gold"])

    doc.add_paragraph()
    add_callout(
        doc,
        "Project Theme",
        "A domain-specific legal research assistant that combines hybrid retrieval, multi-model legal reasoning, citation verification, and human-centered design thinking.",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared for team presentation and project review\nMay 2026")
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.color.rgb = rgb(COLORS["muted"])

    doc.add_page_break()


def add_toc(doc: Document):
    add_h(doc, "Report Navigation", 1)
    add_p(doc, "This document is written for teammates who need to understand the project quickly but also need enough detail to explain it in a review, viva, or design-thinking presentation.")
    rows = [
        ["1", "Project overview", "What OmniLegal is, what problem it solves, and what it does."],
        ["2", "Design research and requirements", "Personas, user research questions, functional needs, and success criteria."],
        ["3", "Design thinking", "Empathy, problem definition, ideation, prototype, testing, and iteration plan."],
        ["4", "Product walkthrough", "Screenshots and examples of the user-facing workflows."],
        ["5", "Technical architecture", "Frontend, backend, RAG pipeline, datasets, models, and verification."],
        ["6", "Feature and API deep dive", "Feature behavior, endpoints, data lifecycle, and implementation files."],
        ["7-12", "Evaluation, testing, risks, future scope, and presentation", "Current test evidence, usability plan, ethics, demo script, and viva Q&A."],
    ]
    add_table(doc, ["Part", "Section", "Why it matters"], rows, [0.55, 1.65, 4.45])


def add_project_overview(doc: Document):
    add_h(doc, "1. Project Overview", 1)
    add_h(doc, "1.1 What the Project Is", 2)
    add_p(doc, "OmniLegal AI is a domain-specific legal research assistant. It is designed to help users ask legal questions, retrieve relevant legal sources, compare jurisdictions, verify citations, and generate structured legal outputs such as IRAC answers, advocacy packets, debate support, and source-backed research explanations.")
    add_p(doc, "The project is not simply a generic chatbot wrapper. Its central idea is verified legal intelligence: the answer should be connected to visible sources, and citations should be checked instead of blindly trusted. The system uses Retrieval-Augmented Generation, known as RAG, where source passages are retrieved first and then used as grounded context for the language model.")
    add_p(doc, "In practical terms, a teammate can explain the project as: a legal AI tool that reads from a curated legal corpus, searches for the most relevant authority, writes an answer in a chosen legal style, and then audits whether the citations are actually supported.")

    add_h(doc, "1.2 What Problem It Solves", 2)
    add_p(doc, "Legal research is difficult because law is long, technical, jurisdiction-sensitive, and citation-heavy. Students and researchers often need to read treaties, constitutions, statutes, court judgments, and commentary before answering one question. General-purpose chatbots may write fluent explanations but can hallucinate legal citations or mix jurisdictions.")
    add_bullets(doc, [
        "Users need faster access to source-backed legal explanations.",
        "Users need clear differences between Indian law, international law, US law, UK law, EU law, and other jurisdictions.",
        "Users need citations that can be checked, not only polished text.",
        "Students and MUN delegates need structured outputs such as issue summaries, debate cards, stance explanations, and IRAC answers.",
        "Teams need an evaluation-aware AI project where results can be tested and discussed honestly.",
    ])

    add_h(doc, "1.3 Core Value Proposition", 2)
    rows = [
        ["Search", "Find relevant legal passages from local and remote corpora.", "Reduces manual reading time."],
        ["Reason", "Generate structured legal analysis using persona and workflow-specific prompts.", "Turns raw sources into understandable answers."],
        ["Compare", "Run cross-jurisdiction analysis such as Indian vs international law.", "Helps explain conflicts, gaps, and similarities."],
        ["Verify", "Check citation markers, quote match, source existence, and support strength.", "Reduces hallucination risk."],
        ["Present", "Show results in React and Chainlit interfaces with source panels and workflow-specific screens.", "Makes the system demonstrable and team-friendly."],
    ]
    add_table(doc, ["Capability", "What it does", "Why it matters"], rows, [1.1, 3.0, 2.35])

    add_h(doc, "1.4 Main Users", 2)
    add_table(doc, ["User", "Need", "How OmniLegal helps"], [
        ["Law student", "Understand legal principles and write exam-ready answers.", "Law Student persona, IRAC format, source-backed explanations."],
        ["Researcher", "Compare legal positions and inspect authorities.", "Researcher persona, citations rail, cross-jurisdiction workflows."],
        ["MUN delegate", "Prepare country stance, speeches, rebuttals, and debate cards.", "Advocacy Studio, conflict detection, brief generation."],
        ["Non-law user", "Understand practical meaning without legal jargon.", "Layman and Tourist personas with simpler explanations."],
        ["Team evaluator", "See whether the project is reliable.", "Benchmark artifacts, citation verification, smoke tests, limitations."],
    ], [1.2, 2.35, 3.0])

    add_h(doc, "1.5 System Capabilities at a Glance", 2)
    add_bullets(doc, [
        "Research Console: persona-based legal question answering with cited source passages.",
        "Conflict Atlas: map-based cross-jurisdiction verdicts for legal topics.",
        "Citation Forensics: audit legal prose and classify citations as verified, partial, suspicious, hallucinated, or not found.",
        "Advocacy Studio: generate position papers, speeches, rebuttals, leverage cards, and argument packs.",
        "Live Authority: query external legal registries such as CourtListener, GovInfo, EUR-Lex, HUDOC, Indian Kanoon, and UN Treaty resources when configured.",
        "Council of Models: compare multiple model answers and synthesize a more reliable conclusion.",
        "Comparative IRAC: produce parallel legal analysis across selected jurisdictions.",
        "Dataset and donor registries: document what corpora, benchmarks, and donor patterns support the system.",
    ])


def add_requirements_and_design_research(doc: Document):
    add_h(doc, "2. Design Research and Requirements", 1)
    add_p(doc, "This section expands the design-thinking evidence behind the project. It can be used directly in a presentation to show that the project was not built only from a technical idea, but from user needs, constraints, and measurable goals.")

    add_h(doc, "2.1 Assumed User Research Questions", 2)
    add_p(doc, "For a classroom design-thinking project, formal interviews may be limited. However, the team can still structure discovery around clear research questions. These questions explain what the team is trying to learn from students, researchers, and MUN users.")
    add_table(doc, ["Research question", "Why it matters", "Design decision influenced"], [
        ["Where do users lose the most time during legal research?", "The project must reduce a real workflow bottleneck.", "Prioritize retrieval, source panels, and ready summaries."],
        ["Which legal outputs do users need most often?", "Different users need different formats.", "Add Research, IRAC, Forensics, Atlas, and Advocacy workflows."],
        ["How do users decide whether to trust an AI answer?", "Trust is central in legal AI.", "Make citations and verification visible."],
        ["What level of language is comfortable for users?", "A beginner and a researcher cannot receive the same answer.", "Add personas and tone control."],
        ["How do users present legal research to others?", "The project is for team presentation and MUN-style use.", "Include screenshots, exportable reports, and debate-ready outputs."],
    ], [2.1, 2.2, 2.25])

    add_h(doc, "2.2 Detailed Personas", 2)
    add_table(doc, ["Persona", "Goals", "Pain points", "Required product behavior"], [
        ["AIDS/legal-tech student", "Understand legal AI, explain architecture, and demonstrate project outcomes.", "May know AI concepts but not legal citation standards.", "Use diagrams, module maps, and simple definitions."],
        ["Law student", "Prepare exam-style or assignment-style legal answers.", "Needs IRAC, leading cases, and reliable source references.", "Provide Law Student persona, IRAC output, and citation rail."],
        ["MUN delegate", "Create speeches, position papers, and rebuttal cards quickly.", "Needs country stance and international law framing under time pressure.", "Provide Advocacy Studio, conflict analysis, and debate support."],
        ["Researcher", "Compare doctrines across jurisdictions and check sources.", "Needs depth, citations, and uncertainty labels.", "Provide Researcher persona, Comparative IRAC, and source panels."],
        ["Non-law user", "Understand legal meaning in plain language.", "Gets confused by jargon and long citations.", "Provide Layman and Tourist modes with simpler wording."],
        ["Evaluator/faculty", "Check project originality, implementation quality, and testing.", "Needs evidence that the system is not only a UI mockup.", "Show architecture, tests, dataset registries, and limitations."],
    ], [1.35, 1.75, 1.75, 1.65])

    add_h(doc, "2.3 Functional Requirements", 2)
    add_table(doc, ["ID", "Requirement", "Implemented through"], [
        ["FR-01", "The system shall accept natural language legal questions.", "Research Console and Chainlit chat."],
        ["FR-02", "The system shall retrieve source passages before answer generation.", "RAG retriever and vector store modules."],
        ["FR-03", "The system shall show citations and source excerpts to users.", "Citation rail and source panels."],
        ["FR-04", "The system shall support different answer personas.", "Tourist, Law Student, Researcher, Layman, Conflict Detector."],
        ["FR-05", "The system shall compare jurisdictions.", "Conflict Atlas and Comparative IRAC."],
        ["FR-06", "The system shall audit legal prose written by users or other AI tools.", "Citation Forensics workflow."],
        ["FR-07", "The system shall generate advocacy or MUN-style material.", "Advocacy Studio and debate support services."],
        ["FR-08", "The system shall expose API endpoints for each major workflow.", "FastAPI routers."],
        ["FR-09", "The system shall maintain dataset and donor registries.", "dataset_registry.json and donor_registry.json."],
        ["FR-10", "The system shall keep evaluation artifacts.", "test_reports and data/evals folders."],
    ], [0.75, 3.1, 2.6])

    add_h(doc, "2.4 Non-Functional Requirements", 2)
    add_table(doc, ["Quality attribute", "Requirement", "How the project addresses it"], [
        ["Reliability", "Answers should not rely only on model memory.", "RAG retrieval and citation verification."],
        ["Explainability", "Users should know why an answer was produced.", "Visible citations, source excerpts, and workflow labels."],
        ["Usability", "Users should choose tasks without prompt engineering.", "Pillar-based React UI and persona tabs."],
        ["Maintainability", "Features should be separable and testable.", "API routers, service modules, pipeline modules, registries."],
        ["Scalability", "The corpus should be extensible.", "Qdrant/vector backend design and ingestion scripts."],
        ["Safety", "The system should avoid pretending to be a lawyer.", "Disclaimers, verification, and limitation sections."],
        ["Testability", "The project should produce measurable evidence.", "Smoke tests, backend tests, completion gates, and evaluation artifacts."],
    ], [1.35, 2.4, 2.7])

    add_h(doc, "2.5 Success Criteria", 2)
    add_bullets(doc, [
        "A teammate can explain the project in one minute without reading the code.",
        "A user can see where legal claims came from.",
        "The interface supports more than one legal workflow, not just generic chat.",
        "The design-thinking process is visible through personas, problem statement, prototypes, screenshots, and testing.",
        "The report honestly separates completed prototype behavior from future production work.",
    ])


def add_design_thinking(doc: Document):
    add_h(doc, "3. Design Thinking Application", 1)
    add_figure(doc, ASSET_DIR / "design_thinking_cycle.png", "Figure 1. Design thinking cycle applied to OmniLegal.", 6.8)
    add_p(doc, "Because this is a design thinking project, the important question is not only whether the software works. The important question is whether the project was shaped around real user problems and tested through usable prototypes. OmniLegal can be explained through the five design thinking stages: empathize, define, ideate, prototype, and test.")

    add_h(doc, "2.1 Empathize", 2)
    add_p(doc, "The empathy stage asks: who struggles with legal research, and what makes the struggle painful? The project targets students, MUN delegates, researchers, and non-law users who need legal clarity but may not have access to expensive legal databases or expert guidance.")
    add_table(doc, ["Observed pain point", "User emotion", "Design response"], [
        ["Long legal documents take too much time to read.", "Overwhelmed", "Use retrieval to surface relevant passages first."],
        ["Legal answers often need exact citations.", "Uncertain", "Show source markers, excerpts, page data, and citation panels."],
        ["Different jurisdictions may disagree.", "Confused", "Build Conflict Atlas and Comparative IRAC views."],
        ["Generic AI answers sound confident even when unsupported.", "Distrust", "Add Citation Forensics and source verification."],
        ["Beginners need simpler wording.", "Intimidated", "Provide personas such as Layman, Tourist, and Law Student."],
    ], [2.0, 1.2, 3.2])
    add_callout(doc, "Empathy insight", "The user does not only want an answer. The user wants confidence that the answer came from a source and is written at the right level of complexity.")

    add_h(doc, "2.2 Define", 2)
    add_p(doc, "The define stage converts empathy findings into a crisp design challenge. For this project, the challenge is: How might we help legal learners and researchers produce fast, understandable, and citation-grounded legal analysis without hiding the sources or the uncertainty?")
    add_bullets(doc, [
        "The system must reduce research time but still preserve evidence.",
        "The system must explain law in different styles for different users.",
        "The system must separate source-grounded claims from AI-inferred claims.",
        "The system must be measurable through citation checks, retrieval metrics, and smoke tests.",
        "The interface must be understandable for a project demo, not only usable by the developer.",
    ])

    add_h(doc, "2.3 Ideate", 2)
    add_p(doc, "The ideation stage generated multiple workflows instead of a single chat screen. This is important because legal work is not one task. A user may ask a question, compare countries, test a citation, prepare a speech, or inspect a doctrine. OmniLegal therefore uses workflow-specific pillars.")
    add_table(doc, ["Idea", "User need", "Final feature"], [
        ["Ask legal questions in different styles", "One answer format does not fit all users.", "Research Console with five personas."],
        ["Map legality by country", "Comparative law is easier to understand visually.", "Conflict Atlas."],
        ["Check another AI's legal answer", "Users need a trust layer.", "Citation Forensics."],
        ["Prepare MUN/legal arguments", "Users need ready-to-present outputs.", "Advocacy Studio and debate support."],
        ["Compare jurisdictions side by side", "Users need structured legal reasoning.", "Comparative IRAC."],
        ["Use multiple models", "Legal reasoning benefits from disagreement checks.", "Council of Models."],
    ], [2.0, 2.3, 2.0])

    add_h(doc, "2.4 Prototype", 2)
    add_p(doc, "The prototype is implemented as a working local application. The frontend is a React application with multiple pages. The backend is a FastAPI service with route groups for health, ingestion, atlas, forensics, advocacy, live search, council, research, reports, graph, doctrine, red-team, and comparative workflows. The Chainlit interface provides a conversational legal research console.")
    add_p(doc, "Prototype decisions were made pragmatically. Heavy legal NLP components are optional so the system can still run in a lightweight mode. Qdrant is the primary vector backend, but fallback behavior exists. External APIs are configurable through environment variables, so the prototype can demonstrate core workflows locally even when every remote integration is not active.")

    add_h(doc, "2.5 Test", 2)
    add_p(doc, "Testing is handled through a mix of backend tests, smoke evaluations, citation checks, ingestion audits, and UI verification. Current artifacts show strong smoke-test citation behavior, while production gates still identify gaps such as RAGAS faithfulness and ingestion metadata quality. This is good for a design-thinking report because it shows iteration rather than pretending the prototype is finished.")
    add_table(doc, ["Test type", "What it checks", "Design lesson"], [
        ["Smoke tests", "Whether the system returns answers without errors.", "Basic prototype viability."],
        ["Citation existence", "Whether cited source markers exist.", "Trust must be measurable."],
        ["Quote match", "Whether quoted text appears in retrieved context.", "Evidence matters more than fluency."],
        ["Comparative backend tests", "Whether comparative and heat-map features work.", "Visual workflows need backend guarantees."],
        ["Production gates", "Whether the system meets release thresholds.", "Academic prototype can be strong but still not production-ready."],
    ], [1.5, 2.65, 2.25])

    add_h(doc, "2.6 Design Thinking Deliverables for Team Presentation", 2)
    add_table(doc, ["Stage", "Artifact to show", "What teammate should say"], [
        ["Empathize", "User personas and pain-point table.", "We studied how different users struggle with legal research and source trust."],
        ["Define", "Problem statement and design criteria.", "The challenge is to provide fast legal answers without losing citation grounding."],
        ["Ideate", "Workflow/pillar map.", "We designed separate workflows for research, comparison, citation audit, and advocacy."],
        ["Prototype", "Screenshots and architecture diagram.", "The idea was converted into a working React + FastAPI + RAG prototype."],
        ["Test", "Evaluation metrics and limitations.", "We tested citations and workflows, then identified improvement areas."],
    ], [1.05, 2.25, 3.15])


def add_screenshots(doc: Document):
    add_h(doc, "4. Product Walkthrough with Screenshots", 1)
    add_p(doc, "The following screenshots are from the local OmniLegal React interface. Exact counts such as chunks and collections can change depending on what has been ingested at the time of running the app.")

    add_h(doc, "3.1 Landing Page", 2)
    add_figure(doc, SCREENSHOT_DIR / "01_landing.png", "Figure 2. Landing page presenting the product promise: verdict, map, and proof.", 6.8)
    add_p(doc, "The landing page communicates the product in one sentence: OmniLegal gives a primary-source verdict, a map, and a forensic audit of citations. The metrics panel shows the current local corpus state. The call-to-action buttons lead users into specific expert workflows instead of a generic blank chat.")

    add_h(doc, "3.2 Feature Sidebar", 2)
    add_figure(doc, SCREENSHOT_DIR / "02_feature_sidebar.png", "Figure 3. Feature navigation sidebar showing main workflows, research tools, and labs.", 6.8)
    add_p(doc, "The sidebar makes the system easier to present because each capability is visible as a named module. This supports design thinking because users can choose a task-oriented workflow: Atlas for comparison, Research for question answering, Forensics for verification, Advocacy for MUN/legal argument generation, and Live for authority lookup.")

    add_h(doc, "3.3 Research Console Example", 2)
    add_figure(doc, SCREENSHOT_DIR / "03_research_erga_omnes.png", "Figure 4. Research Console answering an erga omnes question with citation rail.", 6.8)
    add_p(doc, "This screenshot shows a concrete example: the user asks what erga omnes obligations each country has. The interface includes persona tabs, a query input, generated legal explanation, model metadata, and a citations rail. The main design benefit is that the answer and the evidence are visible together.")
    add_callout(doc, "Example explanation for presentation", "A teammate can say: Here the user asks a legal concept question. OmniLegal returns a source-backed answer and places the citations on the right, so the user can inspect where the answer came from.")

    add_h(doc, "3.4 Conflict Atlas Example", 2)
    add_figure(doc, SCREENSHOT_DIR / "04_atlas.png", "Figure 5. Conflict Atlas input screen for cross-jurisdiction legal mapping.", 6.8)
    add_p(doc, "The Conflict Atlas is designed for comparative law topics. A user types a legal topic such as right to self-determination, death penalty for drug trafficking, hate speech laws online, or surveillance without warrant. The system retrieves legal sources for supported jurisdictions and can mark fallback AI-inferred results separately.")

    add_h(doc, "3.5 Comparative IRAC Example", 2)
    add_figure(doc, SCREENSHOT_DIR / "05_comparative.png", "Figure 6. Comparative IRAC page for side-by-side jurisdiction analysis.", 6.8)
    add_p(doc, "The Comparative IRAC workflow is useful for academic explanation because IRAC is familiar in legal education: Issue, Rule, Application, and Conclusion. The user selects jurisdictions and the system produces parallel legal analysis, making agreements and conflicts easier to discuss.")


def add_architecture(doc: Document):
    add_h(doc, "5. Technical Architecture", 1)
    add_figure(doc, ASSET_DIR / "technical_flow.png", "Figure 7. High-level technical flow from user query to verified answer.", 6.9)
    add_h(doc, "4.1 Repository Structure", 2)
    add_table(doc, ["Folder/file", "Role in the project"], [
        ["frontend/", "React web interface with pages for Atlas, Research, Comparative, Forensics, Advocacy, Live, Council, Library, and other workflows."],
        ["backend/server.py", "FastAPI application host that mounts multiple API routers and exposes /api routes."],
        ["src/api_router*.py", "Route definitions for health, ingestion, atlas, forensics, advocacy, live authority, council, research, reports, comparative, and SOTA workflows."],
        ["src/pipeline/", "Main legal reasoning graph: classification, entity extraction, retrieval, jurisdiction analysis, synthesis, source gate, and citation verification."],
        ["pipeline_v2/", "Newer RAG ingestion and retrieval modules including retriever, reranker, vector store, seed corpus, and orchestrator."],
        ["data/", "Dataset registry, donor registry, ingestion summary, evaluation artifacts, and gold datasets."],
        ["caselaws/", "Jurisdiction-specific case-law configuration and source groupings."],
        ["tests/ and backend/tests/", "Regression tests and workflow-specific backend tests."],
        ["reports/", "Generated report documents, screenshots, and render QA artifacts."],
    ], [1.75, 4.9])

    add_h(doc, "4.2 Layered System Design", 2)
    add_table(doc, ["Layer", "Main responsibility", "Examples"], [
        ["User Interface", "Collect user intent and show outputs clearly.", "React pages, Chainlit chat, personas, citation panels, map views."],
        ["API Layer", "Expose stable workflow endpoints.", "/api/research/ask, /api/atlas/analyze, /api/forensics/verify, /api/compare/analyze."],
        ["Pipeline Layer", "Coordinate classification, retrieval, reasoning, synthesis, and verification.", "src/pipeline/graph.py, synthesizer.py, citation_verifier.py."],
        ["Retrieval Layer", "Find relevant source passages.", "Dense retrieval, sparse retrieval, RRF fusion, reranking, Qdrant or fallback."],
        ["Corpus Layer", "Store legal documents, case-law metadata, datasets, and source registries.", "Indian Constitution, UN Charter, ICCPR, ICESCR, Shaw commentary, case-law configs."],
        ["Evaluation Layer", "Measure reliability and identify gaps.", "Smoke tests, citation existence, quote match, RAGAS, production gates."],
    ], [1.25, 2.7, 2.55])

    add_h(doc, "4.3 RAG Pipeline Explained Simply", 2)
    add_numbered(doc, [
        "The user enters a legal question or selects a workflow.",
        "The backend classifies the query and extracts legal entities, issues, jurisdictions, and citation hints.",
        "The retriever searches indexed legal corpora for relevant passages.",
        "A reranker can reorder passages so the strongest legal evidence appears first.",
        "The LLM generates an answer using retrieved passages as context.",
        "The verifier checks whether citations and source claims are supported.",
        "The frontend displays the final answer with source panels, trust labels, and workflow-specific views.",
    ])

    add_h(doc, "4.4 Models and AI Components", 2)
    add_table(doc, ["Component", "Example configuration", "Purpose"], [
        ["Generation model", "Groq-hosted Llama 3.3 70B, with optional Gemini/Groq/local fallbacks depending on environment.", "Generate legal explanations, summaries, briefs, and comparative analysis."],
        ["Embedding model", "BAAI/bge-m3.", "Convert text chunks and queries into vectors for semantic retrieval."],
        ["Reranker", "BAAI/bge-reranker-v2-m3.", "Improve retrieval precision by scoring candidate passages."],
        ["Classification", "Zero-shot and heuristic classifiers.", "Route legal issues and workflows."],
        ["NER/entity extraction", "spaCy legal NER, GLiNER, and fallback heuristics.", "Find statutes, cases, countries, organizations, dates, and legal concepts."],
        ["Verification", "Lexical support checks and optional NLI/hallucination evaluation.", "Reduce unsupported legal claims."],
    ], [1.6, 2.4, 2.6])

    add_h(doc, "4.5 Data and Corpus Design", 2)
    add_p(doc, "The dataset registry currently lists 17 datasets or dataset references. These include runtime corpora, training datasets, evaluation-only benchmarks, manual gold sets, and donor-inspired legal NLP patterns. The donor registry lists 9 donors or references such as Law-AI summarization, CLERC, LLeQA, BSARD, Mining Legal Arguments, FairLex, CaseHOLD, and LexGLUE.")
    add_table(doc, ["Source group", "Use in project"], [
        ["Local PDF corpus", "Runtime retrieval and evaluation using legal texts such as Indian Constitution, UN Charter, ICCPR, ICESCR, and international law commentary."],
        ["Curated legal NLP datasets", "Training, evaluation, retrieval QA patterns, summarization patterns, and stance prediction support."],
        ["Manual gold sets", "Conflict detection, stance prediction, and brief-review evaluation."],
        ["Case-law registries", "Jurisdiction-specific source coverage for India, US, UK, EU, Russia, Israel, and international sources."],
        ["Remote adapters", "Potential live ingestion from official or legal information sources when credentials and licensing permit."],
    ], [2.0, 4.45])

    add_h(doc, "4.6 Why the Architecture Fits the Problem", 2)
    add_p(doc, "Legal AI needs more structure than a normal chatbot. A normal chatbot can produce attractive text, but legal research needs traceability, jurisdiction labels, uncertainty handling, and source verification. OmniLegal separates these responsibilities into layers so that each part can be tested and improved.")
    add_bullets(doc, [
        "Separate UI workflows reduce cognitive load for users.",
        "Separate API routers make features easier to test and explain.",
        "Separate retrieval and generation prevent the model from relying only on memory.",
        "Separate citation verification creates a trust layer after generation.",
        "Separate dataset registries make the corpus auditable for academic reporting.",
    ])


def add_feature_and_api_deep_dive(doc: Document):
    add_h(doc, "6. Feature and API Deep Dive", 1)
    add_h(doc, "6.1 Feature Deep Dive", 2)
    add_table(doc, ["Feature", "Input", "Processing", "Output"], [
        ["Research Console", "Legal question and persona.", "Retrieve passages, generate answer, verify citations.", "Persona-specific cited answer with citations rail."],
        ["Conflict Atlas", "Legal topic and AI fallback option.", "Retrieve jurisdiction-specific authorities and infer unsupported countries cautiously.", "Legal/restricted/illegal/no-data map and country details."],
        ["Citation Forensics", "Pasted legal prose.", "Extract citations, search corpus, compare support, grade claims.", "Annotated text and claim-by-claim trust report."],
        ["Advocacy Studio", "Country, issue, and position.", "Retrieve authorities, synthesize stance, create argument material.", "Position paper, speech, rebuttals, and leverage cards."],
        ["Live Authority", "Search query and source choices.", "Call configured external legal registries.", "Live results from primary or legal-information sources."],
        ["Council of Models", "Legal query.", "Ask multiple models and compare disagreement.", "Consensus answer or resolved verdict."],
        ["Comparative IRAC", "Question and selected jurisdictions.", "Run parallel Issue, Rule, Application, Conclusion analysis.", "Jurisdiction cards, heat map, synthesis, and gaps."],
        ["Library", "Search/filter terms.", "Inspect indexed legal chunks and collections.", "Corpus transparency and source exploration."],
    ], [1.55, 1.55, 2.05, 1.95])

    add_h(doc, "6.2 API Catalogue", 2)
    add_p(doc, "The frontend communicates with the backend through a consistent /api contract. This makes the prototype easier to test because every visible workflow maps to an endpoint.")
    add_table(doc, ["Endpoint", "Method", "Purpose"], [
        ["/api/health", "GET", "Check backend availability."],
        ["/api/overview", "GET", "Return corpus and model overview for the landing page."],
        ["/api/ingestion/status", "GET", "Show ingestion health and collection status."],
        ["/api/atlas/analyze", "POST", "Run cross-jurisdiction conflict atlas analysis."],
        ["/api/forensics/verify", "POST", "Audit citations in pasted legal prose."],
        ["/api/advocacy/generate", "POST", "Generate advocacy or MUN-oriented outputs."],
        ["/api/live/search", "POST", "Search live external legal authority sources."],
        ["/api/council/debate", "POST", "Run multi-model council reasoning."],
        ["/api/research/ask", "POST", "Answer legal questions with persona and citations."],
        ["/api/debug/retrieve", "GET", "Inspect retrieval results for a query."],
        ["/api/compare/analyze", "POST", "Run Comparative IRAC."],
        ["/api/compare/longitudinal", "POST", "Run longitudinal heat map analysis where configured."],
    ], [2.05, 0.85, 3.55])

    add_h(doc, "6.3 Data Lifecycle", 2)
    add_numbered(doc, [
        "Source selection: legal documents, public legal datasets, case-law registries, and manual gold sets are selected.",
        "Preprocessing: PDFs and text files are parsed, cleaned, chunked, and assigned metadata.",
        "Indexing: chunks are embedded and stored in a vector backend such as Qdrant or a fallback store.",
        "Retrieval: user queries are transformed into search requests against the indexed corpus.",
        "Reranking: candidate passages are reordered to improve relevance.",
        "Generation: the LLM writes an answer based on retrieved context.",
        "Verification: citations, quotes, and claim support are checked.",
        "Presentation: the UI shows the answer, evidence, trust signals, and workflow-specific results.",
        "Evaluation: metrics and test artifacts are stored for future improvement.",
    ])

    add_h(doc, "6.4 Important Implementation Files", 2)
    add_table(doc, ["File/module", "Why teammates should know it"], [
        ["frontend/src/App.js", "Shows the main routes and feature pages in the React application."],
        ["frontend/src/lib/api.js", "Shows how frontend workflows call backend endpoints."],
        ["frontend/src/pages/Research.js", "Implements persona-based research UI."],
        ["frontend/src/pages/Atlas.js", "Implements Conflict Atlas UI and map interaction."],
        ["frontend/src/pages/Forensics.js", "Implements citation audit UI."],
        ["backend/server.py", "Bootstraps the FastAPI app and mounts routers."],
        ["src/api_router_v2.py", "Contains major API pillars such as Atlas, Forensics, Advocacy, Live, Council, Research."],
        ["src/pipeline/graph.py", "Coordinates the legal reasoning pipeline."],
        ["src/pipeline/citation_verifier.py", "Checks generated citations and source support."],
        ["pipeline_v2/retriever.py", "Handles retrieval logic for RAG."],
        ["pipeline_v2/reranker.py", "Improves ranking of retrieved passages."],
        ["data/dataset_registry.json", "Documents datasets and their use cases."],
        ["data/donor_registry.json", "Documents external donor patterns and benchmark references."],
    ], [2.25, 4.2])


def add_examples(doc: Document):
    add_h(doc, "7. Example Use Cases", 1)
    add_h(doc, "5.1 Example 1: Research Question", 2)
    add_p(doc, "User question: What erga omnes obligations does each country have under international law?")
    add_table(doc, ["Step", "What happens", "User sees"], [
        ["Input", "The user selects Researcher persona and asks the question.", "A query box and persona tabs."],
        ["Retrieve", "The system searches for international law authorities and relevant commentary.", "Citations such as ICJ cases and explanatory excerpts."],
        ["Generate", "The LLM writes an academic explanation using retrieved passages.", "A paragraph-style answer with citation markers."],
        ["Verify", "Citation support is checked.", "Trust badge and citation rail."],
        ["Learn", "The user reads the answer and expands citations.", "A source-backed explanation instead of unsupported prose."],
    ], [1.1, 3.0, 2.35])

    add_h(doc, "5.2 Example 2: Citation Forensics", 2)
    add_p(doc, "User action: Paste a paragraph written by another AI that contains legal citations.")
    add_p(doc, "Expected output: The Forensics workflow extracts citations, retrieves best matches, calculates overlap, and marks each claim as verified, partial, suspicious, hallucinated, not found, or no citation. This is especially useful for design thinking because it addresses the user's trust problem directly.")

    add_h(doc, "5.3 Example 3: MUN or Debate Preparation", 2)
    add_p(doc, "User action: Choose a country, topic, and position in Advocacy Studio.")
    add_bullets(doc, [
        "The system retrieves supporting legal sources.",
        "It generates a position paper and opening speech.",
        "It adds rebuttal cards and leverage cards.",
        "It can surface conflicts between domestic and international positions.",
        "The result is closer to a presentation artifact than a normal chatbot answer.",
    ])

    add_h(doc, "5.4 Example 4: Comparative IRAC", 2)
    add_p(doc, "User question: Compare the right to privacy under Indian, US, and UK constitutional law.")
    add_table(doc, ["IRAC part", "Meaning", "How OmniLegal uses it"], [
        ["Issue", "The legal question being answered.", "Identifies the exact legal conflict or comparison."],
        ["Rule", "The legal rule from statutes, cases, treaties, or doctrine.", "Retrieves and cites jurisdiction-specific authorities."],
        ["Application", "How the rule applies to the facts or topic.", "Generates jurisdiction-specific analysis."],
        ["Conclusion", "The answer or outcome.", "Summarizes similarities, conflicts, and gaps."],
    ], [1.2, 2.35, 2.9])


def add_evaluation(doc: Document):
    add_h(doc, "8. Evaluation and Current Results", 1)
    add_p(doc, "The report should present OmniLegal as a strong evaluated academic prototype, not as a finished professional legal advice product. This distinction is important and responsible.")
    add_table(doc, ["Metric/artifact", "Reported value or status", "Interpretation"], [
        ["Latest legal smoke run", "54 queries, 0 errors.", "The pipeline can handle a non-trivial set of local smoke queries."],
        ["Hallucination rate", "0.0000 in the referenced smoke artifact.", "Strong controlled-test behavior, but not a universal guarantee."],
        ["Citation existence rate", "1.0000 in the referenced smoke artifact.", "Citations existed for the generated claims in that test."],
        ["Quote match", "1.0000 in the referenced smoke artifact.", "Quoted text matched retrieved material in that test."],
        ["Total citations", "204 total, 204 correct in the referenced smoke artifact.", "Good evidence that the verifier and retrieval path are useful."],
        ["Retrieval smoke recall", "Recall@5 and Recall@10 reported as 0.3333 in older report artifacts.", "Retrieval still needs stronger datasets and tuning."],
        ["Comparative IRAC tests", "13/13 backend tests passed in iteration_7.json.", "Comparative heat map and query expansion were validated locally."],
        ["Production gate", "not_ready.", "The project should be described as an academic prototype with clear improvement work."],
        ["RAGAS faithfulness", "Gate not fully satisfied in completion artifacts.", "Faithfulness evaluation needs refinement and reruns."],
        ["Ingestion metadata quality", "Known gaps in document hash, canonical ID, legal type, and importance score.", "Metadata should be improved before serious deployment."],
    ], [1.85, 1.85, 2.8])

    add_h(doc, "6.1 What the Results Mean", 2)
    add_p(doc, "The strongest current result is that the system is demonstrable and evaluation-aware. It has real UI screens, route-level architecture, dataset registries, source verification, and test artifacts. The results support an academic claim that the team built a working prototype for legal research assistance.")
    add_p(doc, "The limitations are equally important. Legal AI requires high accuracy, source licensing clarity, strong metadata, and expert review. The project should not be presented as ready for professional legal advice. It is better to present it as a responsibly tested prototype with a clear path to improvement.")

    add_h(doc, "6.2 Suggested Next Tests", 2)
    add_bullets(doc, [
        "Expand gold datasets for legal QA, conflict detection, stance prediction, and brief generation.",
        "Run retrieval evaluation by jurisdiction and document type.",
        "Add human review labels for whether answers are legally correct, not only citation-present.",
        "Improve metadata fields for every ingested chunk.",
        "Measure user task completion time before and after using OmniLegal.",
        "Run usability testing with teammates acting as student, researcher, MUN delegate, and layman personas.",
    ])


def add_testing_and_usability_plan(doc: Document):
    add_h(doc, "9. Usability Testing Plan", 1)
    add_p(doc, "For design thinking, technical tests are not enough. The team should also test whether people can actually understand and use the prototype. The following plan can be used for classroom demonstration or future evaluation.")
    add_h(doc, "9.1 Usability Test Tasks", 2)
    add_table(doc, ["Task", "Participant instruction", "Success indicator"], [
        ["Research task", "Ask a legal question and identify one source used by the answer.", "Participant can find answer and citation rail."],
        ["Forensics task", "Paste a sample paragraph and explain which citation is risky.", "Participant can interpret trust labels."],
        ["Atlas task", "Enter a legal topic and explain what AI-inferred fallback means.", "Participant understands grounded vs inferred results."],
        ["Comparative task", "Select three jurisdictions and explain one legal difference.", "Participant can read IRAC comparison."],
        ["Presentation task", "Use the report to explain OmniLegal in one minute.", "Participant can state problem, solution, architecture, and testing."],
    ], [1.35, 3.0, 2.1])

    add_h(doc, "9.2 Usability Metrics", 2)
    add_table(doc, ["Metric", "How to measure", "Good target"], [
        ["Task completion", "Did the participant finish the assigned task?", "80 percent or higher in classroom testing."],
        ["Time to source", "How long until the participant finds a citation/source?", "Under 60 seconds for Research Console."],
        ["Trust understanding", "Can participant explain verified vs partial vs hallucinated?", "Correct explanation after one walkthrough."],
        ["Navigation clarity", "Can participant choose the right workflow from sidebar?", "Minimal prompting needed."],
        ["Presentation clarity", "Can teammate explain the project without code?", "Clear explanation of users, problem, RAG, and verification."],
    ], [1.5, 2.65, 2.3])

    add_h(doc, "9.3 Feedback Form Questions", 2)
    add_bullets(doc, [
        "What was the easiest part of the interface to understand?",
        "Where did you feel confused?",
        "Did the citations make the answer more trustworthy? Why or why not?",
        "Which workflow would you use most: Research, Forensics, Atlas, Advocacy, or Comparative IRAC?",
        "What feature should be improved before public deployment?",
        "Could you explain the project to another teammate after using the report?",
    ])

    add_h(doc, "9.4 Acceptance Criteria for the Next Iteration", 2)
    add_bullets(doc, [
        "Every major screen should have a sample query or example button.",
        "Every generated answer should show visible source evidence or a clear warning when evidence is missing.",
        "Every AI-inferred result should be visually different from source-grounded output.",
        "Every project demo should include at least one citation verification example.",
        "Every future report should include the exact test date and artifact path for evaluation metrics.",
    ])


def add_risks_future(doc: Document):
    add_h(doc, "10. Ethics, Risks, and Limitations", 1)
    add_table(doc, ["Risk", "Why it matters", "Mitigation"], [
        ["Legal advice risk", "Users may overtrust generated legal explanations.", "Clear disclaimer: research assistance only, verify with primary sources and experts."],
        ["Hallucinated citations", "Wrong legal citations can mislead users.", "Citation Forensics, quote matching, citation panels, and fallback grounded drafts."],
        ["Jurisdiction confusion", "A rule from one country may not apply in another.", "Jurisdiction labels, Comparative IRAC, Atlas, and explicit conflict notes."],
        ["Outdated law", "Law changes over time.", "Live Authority sources, update logs, and future doctrine drift tracking."],
        ["Licensing", "Some legal sources may have redistribution limits.", "Registry license notes and source approval before public release."],
        ["Bias and fairness", "Legal data may reflect historical bias.", "FairLex-style evaluation references and human review."],
        ["Privacy", "Uploaded legal documents may contain sensitive data.", "Future user workspace controls, redaction, and privacy policy."],
    ], [1.45, 2.35, 2.65])

    add_h(doc, "11. Future Scope", 1)
    add_bullets(doc, [
        "Increase corpus coverage with more official legal sources and better jurisdiction metadata.",
        "Create a larger human-reviewed gold dataset for legal QA and comparative law.",
        "Fine-tune or train a reranker on legal retrieval triples.",
        "Add stronger multilingual support for Indian regional languages and international legal sources.",
        "Add authenticated user accounts and private document workspaces.",
        "Add export to DOCX and PDF for generated briefs, position papers, and citation audits.",
        "Create visual dashboards for corpus health, source coverage, and evaluation trends.",
        "Add expert-in-the-loop review for high-risk legal answers.",
        "Deploy with monitored Qdrant, secret management, rate limiting, logs, and safety controls.",
        "Improve RAGAS faithfulness and production gates until the system reaches a release-ready threshold.",
    ])


def add_presentation(doc: Document):
    add_h(doc, "12. How to Present This Project to the Team", 1)
    add_h(doc, "9.1 One-Minute Explanation", 2)
    add_callout(
        doc,
        "Short presentation script",
        "OmniLegal AI is a legal research assistant built with design thinking. We started from the pain point that legal research is slow and AI answers are hard to trust. Our solution combines a React interface, FastAPI backend, retrieval-augmented generation, legal datasets, and citation verification. Users can ask legal questions, compare jurisdictions, verify citations, and generate debate-ready legal outputs. The important design principle is that every answer should be explainable through visible sources.",
    )

    add_h(doc, "9.2 Suggested Slide Order", 2)
    add_numbered(doc, [
        "Title: OmniLegal AI - Verified Legal Intelligence.",
        "Problem: legal research is slow, complex, and citation-sensitive.",
        "Users: students, researchers, MUN delegates, non-law users, evaluators.",
        "Design thinking process: empathize, define, ideate, prototype, test.",
        "Product screenshots: landing page, sidebar, research console, atlas, comparative IRAC.",
        "Architecture: React + FastAPI + RAG pipeline + corpus + verifier.",
        "Datasets and sources: local legal PDFs, registries, donor datasets, gold sets.",
        "Example walkthrough: ask a question and inspect citations.",
        "Evaluation: smoke results, citation checks, comparative tests, current limitations.",
        "Future scope and responsible AI conclusion.",
    ])

    add_h(doc, "9.3 Demo Script", 2)
    add_table(doc, ["Demo step", "What to show", "What to say"], [
        ["Landing", "Open the landing screenshot or app home.", "This is the product promise: verdict, map, proof."],
        ["Research", "Show the erga omnes screenshot.", "The answer is not alone; citations are visible beside it."],
        ["Atlas", "Show the topic input.", "Comparative law becomes visual and easier to discuss."],
        ["Comparative IRAC", "Show jurisdiction selection.", "Legal analysis is structured by jurisdiction and IRAC."],
        ["Architecture", "Show the technical flow diagram.", "The model retrieves sources before generating and verifies after generating."],
        ["Testing", "Show evaluation table.", "We tested the prototype and identified what still needs improvement."],
    ], [1.3, 2.2, 2.95])

    add_h(doc, "9.4 Teammate Speaking Roles", 2)
    add_table(doc, ["Speaker", "Section", "Focus"], [
        ["Member 1", "Problem and empathy", "Explain users, pain points, and why legal AI must be source-grounded."],
        ["Member 2", "Design thinking", "Explain define, ideate, prototype, and testing decisions."],
        ["Member 3", "System architecture", "Explain frontend, backend, RAG, datasets, and models."],
        ["Member 4", "Demo and screenshots", "Walk through Research, Atlas, Forensics, and Comparative IRAC."],
        ["Member 5", "Evaluation and future work", "Explain metrics, limitations, ethics, and next improvements."],
    ], [1.0, 1.8, 3.65])

    add_h(doc, "12.5 Common Questions and Suggested Answers", 2)
    add_table(doc, ["Question", "Suggested answer"], [
        ["Is OmniLegal just ChatGPT for law?", "No. It is workflow-based and source-grounded. It retrieves legal passages, generates an answer, and verifies citations."],
        ["Why is this a design thinking project?", "The system was structured around user pain points: legal complexity, lack of trust, different expertise levels, and presentation needs."],
        ["What is the main innovation?", "Combining legal RAG, citation verification, comparative workflows, and persona-based outputs in one academic prototype."],
        ["Can it replace a lawyer?", "No. It is a research assistant and educational prototype. Legal outputs must be verified by qualified experts."],
        ["What makes the project testable?", "The codebase includes backend tests, smoke artifacts, citation metrics, completion gates, and evaluation registries."],
        ["What is the biggest current limitation?", "Production readiness is not complete. Retrieval recall, metadata quality, RAGAS faithfulness, and source licensing need improvement."],
        ["How can teammates contribute?", "One can improve UI, one can expand datasets, one can write tests, one can handle evaluation, and one can prepare presentation/demo material."],
    ], [2.2, 4.25])

    add_h(doc, "12.6 Suggested Project Timeline", 2)
    add_table(doc, ["Phase", "Activities", "Deliverables"], [
        ["Week 1: Empathize", "Identify users, collect pain points, inspect existing legal research workflows.", "Personas, empathy map, user journey."],
        ["Week 2: Define", "Write problem statement, success criteria, and design constraints.", "Problem statement and requirements."],
        ["Week 3: Ideate", "Choose workflows and screen concepts.", "Feature map and prototype plan."],
        ["Week 4-5: Prototype", "Build React screens, backend routes, retrieval pipeline, and verification flow.", "Working local prototype."],
        ["Week 6: Test", "Run backend tests, smoke tests, UI walkthroughs, and team feedback sessions.", "Evaluation table and improvement backlog."],
        ["Week 7: Present", "Prepare screenshots, report, demo script, and role distribution.", "Final report and presentation deck."],
    ], [1.3, 3.0, 2.15])


def add_appendices(doc: Document):
    add_h(doc, "Appendix A: Glossary", 1)
    add_table(doc, ["Term", "Meaning"], [
        ["RAG", "Retrieval-Augmented Generation. The system retrieves documents first and uses them as context for generation."],
        ["IRAC", "Issue, Rule, Application, Conclusion. A standard legal reasoning structure."],
        ["Citation verification", "Checking whether citations and quoted/supporting text exist in retrieved sources."],
        ["Embedding", "A numeric representation of text used for semantic search."],
        ["Vector database", "A database optimized for searching embeddings; Qdrant is used as the primary vector backend."],
        ["Reranker", "A model that reorders retrieved passages by relevance."],
        ["Corpus", "The collection of documents available for retrieval."],
        ["Persona", "A user-facing answer mode that changes tone, structure, and depth."],
        ["Gold dataset", "A human-created or curated evaluation set used to measure system performance."],
    ], [1.5, 4.95])

    add_h(doc, "Appendix B: Design Thinking Canvas", 1)
    add_table(doc, ["Canvas item", "OmniLegal answer"], [
        ["Who are we designing for?", "Students, researchers, MUN delegates, non-law users, and evaluators who need legal clarity with evidence."],
        ["What do they need?", "Fast legal explanations, source trust, jurisdiction comparison, and usable presentation outputs."],
        ["What is the main frustration?", "Legal documents are long and AI answers can be unsupported."],
        ["What is the design challenge?", "Help users produce citation-grounded legal analysis without hiding uncertainty."],
        ["What is the prototype?", "React + Chainlit UI, FastAPI backend, RAG pipeline, corpus, and citation verifier."],
        ["How do we test it?", "Smoke tests, citation metrics, UI walkthroughs, comparative tests, user feedback, and production gates."],
        ["What counts as success?", "Users understand answers faster and can inspect the supporting sources."],
    ], [1.85, 4.6])

    add_h(doc, "Appendix C: Feature-to-Need Matrix", 1)
    add_table(doc, ["Feature", "Need addressed", "Presentation example"], [
        ["Research Console", "Ask legal questions with different depth levels.", "Erga omnes obligations question."],
        ["Citation Forensics", "Check whether legal text is trustworthy.", "Paste AI-generated legal paragraph and audit citations."],
        ["Conflict Atlas", "Understand cross-country legal differences.", "Right to self-determination or surveillance without warrant."],
        ["Comparative IRAC", "Compare legal rules across jurisdictions.", "Right to privacy in India, US, and UK."],
        ["Advocacy Studio", "Generate debate and MUN-ready outputs.", "Country position paper and rebuttal cards."],
        ["Live Authority", "Search external legal authorities.", "CourtListener, GovInfo, EUR-Lex, HUDOC, UN treaty sources."],
        ["Council of Models", "Reduce single-model weakness.", "Compare answers and synthesize consensus."],
    ], [1.6, 2.65, 2.2])


def add_header_footer(doc: Document):
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "OmniLegal AI | Detailed Design Thinking Project Report"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.color.rgb = rgb(COLORS["muted"])
        footer = section.footer.paragraphs[0]
        footer.text = "Prepared for team presentation | Academic prototype, not legal advice"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.color.rgb = rgb(COLORS["muted"])


def main():
    make_design_cycle(ASSET_DIR / "design_thinking_cycle.png")
    make_flow_diagram(ASSET_DIR / "technical_flow.png")
    make_box_diagram(
        ASSET_DIR / "user_journey.png",
        "User Journey: From Confusion to Source-Backed Answer",
        "The product design follows the user from a vague legal question to a verified, explainable output.",
        [
            ("1. Ask", "User enters a legal question, topic, claim, or clause."),
            ("2. Route", "System chooses the right workflow: research, atlas, forensics, advocacy, or IRAC."),
            ("3. Retrieve", "Relevant legal passages are pulled from corpus and registries."),
            ("4. Reason", "LLM generates answer using retrieved evidence and selected persona."),
            ("5. Verify", "Citations and claims are checked before presentation."),
            ("6. Present", "User sees answer, citations, trust labels, and next action."),
        ],
        "Design principle: every screen should reduce uncertainty, not only produce text.",
    )

    doc = Document()
    set_styles(doc)
    add_cover(doc)
    add_toc(doc)
    add_project_overview(doc)
    add_requirements_and_design_research(doc)
    add_design_thinking(doc)
    add_figure(doc, ASSET_DIR / "user_journey.png", "Figure 8. User journey from vague question to verified legal output.", 6.8)
    add_screenshots(doc)
    add_architecture(doc)
    add_feature_and_api_deep_dive(doc)
    add_examples(doc)
    add_evaluation(doc)
    add_testing_and_usability_plan(doc)
    add_risks_future(doc)
    add_presentation(doc)
    add_appendices(doc)
    add_header_footer(doc)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
